from fastapi import FastAPI, Request, Query
from fastapi.openapi.utils import get_openapi
from fastapi.responses import JSONResponse
from docx import Document
from io import BytesIO

app = FastAPI()

def custom_openapi():
    if app.openapi_schema:
        return app.openapi_schema
    openapi_schema = get_openapi(
        title="My Custom API",
        version="1.0.0",
        description="This is a custom API for demonstration purposes",
        routes=app.routes,
    )
    app.openapi_schema = openapi_schema
    return app.openapi_schema

@app.get("/docs", include_in_schema=False)
async def get_docs(request: Request):
    openapi_schema = custom_openapi()

    # Create a new Word document
    doc = Document()

    # Add title and description to the document
    doc.add_heading(openapi_schema["info"]["title"], level=1)
    doc.add_paragraph(openapi_schema["info"]["description"])

    # Add paths and operations to the document
    for path, path_item in openapi_schema["paths"].items():
        doc.add_heading(path, level=2)
        for method, operation in path_item.items():
            doc.add_heading(method.upper(), level=3)
            doc.add_paragraph(operation["summary"])
            doc.add_paragraph(operation["description"])

    # Save the document to a BytesIO object
    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)

    # Set the appropriate headers for downloading the Word document
    headers = {
        "Content-Disposition": "attachment; filename=openapi.docx",
        "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }

    return Response(content=doc_stream, headers=headers)

@app.get("/openapi.json", include_in_schema=False)
async def get_openapi_json():
    openapi_schema = custom_openapi()
    return JSONResponse(content=openapi_schema)




# 가상의 수업 데이터
courses = [
    {
        "id": "CS201",
        "과목명": "데이타베이스 구조",
        "담당교수": "홍길동",
        "구분" : "전공필수",
        "학점": 3,
        "학과" : "전산학과",
        "요일" : ["월","수"],
        "시간" : "10:00 ~ 12:00"
    },
    {
        "id": "CS202",
        "과목명": "이산구조",
        "담당교수": "김갑돌",
        "구분" : "전공필수",
        "학점": 3,
        "학과" : "전산학과",
        "요일" : ["화","목"],
        "시간" : "10:00 ~ 12:00"
    },
    {
        "id": "CS101",
        "과목명": "프로그래밍의 이해",
        "담당교수": "김철수",
        "구분" : "전공선택",
        "학점": 3,
        "학과" : "전산학과",
        "요일" : ["월","수"],
        "시간" : "14:00 ~ 16:00"
    },
    {
        "id": "CS301",
        "과목명": "프로그래밍 언어",
        "담당교수": "이영희",
        "구분" : "전공선택",
        "학점": 3,
        "학과" : "전산학과",
        "요일" : ["월","수"],
        "시간" : "14:00 ~ 16:00"
    }
]

@app.get("/courses")
async def get_courses(
    id: int = Query(default=None),
    과목명: str = Query(default=None),
    담당교수: str = Query(default=None),
    학점: int = Query(default=None, gt=0),
    최소학점: int = Query(default=None, gt=0),
    학과: str = Query(default=None),
    구분: str = Query(default=None)
):
    filtered_courses = courses

    if id is not None:
        # id 필터링
        filtered_courses = [course for course in filtered_courses if course["id"] == id]

    if 과목명 is not None:
        # 이름 포함 필터링
        filtered_courses = [course for course in filtered_courses if 과목명.lower() in course["과목명"].lower()]

    if 담당교수 is not None:
        # 교수명 포함 필터링
        filtered_courses = [course for course in filtered_courses if 담당교수.lower() in course["담당교수"].lower()]
        
    if 구분 is not None:
        # 구분 포함 필터링
        filtered_courses = [course for course in filtered_courses if 구분.lower() in course["구분"].lower()]
        
    if 학과 is not None:
        # 학과 포함 필터링
        filtered_courses = [course for course in filtered_courses if 학과.lower() in course["학과"].lower()]

    if 학점 is not None:
        # 학점 필터링
        filtered_courses = [course for course in filtered_courses if course["학점"] >= credits]

    if 최소학점 is not None:
        # 최소 학점 필터링
        filtered_courses = [course for course in filtered_courses if course["최소학점"] >= min_credits]

    return filtered_courses


